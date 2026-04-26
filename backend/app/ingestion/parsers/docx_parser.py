import io

from docx import Document

from app.ingestion.parsers.base import BaseParser, ParsedDocument

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class DOCXParser(BaseParser):
    @classmethod
    def can_parse(cls, filename: str, mime_type: str = "") -> bool:
        return filename.lower().endswith(".docx") or mime_type == _DOCX_MIME

    async def parse(self, content: bytes, filename: str, **kwargs) -> ParsedDocument:
        doc = Document(io.BytesIO(content))

        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        author = None
        try:
            author = doc.core_properties.author or None
        except Exception:
            pass

        return ParsedDocument(
            raw_text="\n\n".join(paragraphs),
            paragraphs=paragraphs,
            metadata={
                "filename": filename,
                "source_type": "docx",
                "page_count": None,
                "author": author,
                "file_size_bytes": len(content),
            },
        )
